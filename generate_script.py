import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_script():
    doc = docx.Document()

    # Title
    title = doc.add_heading('Mamba Phase 2 Progress Update - Video Script', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('Speaker: David').bold = True
    doc.add_paragraph('Format: Direct to camera / Screen recording walkthrough').italic = True
    doc.add_paragraph()

    def add_section(title_text, script_text, ui_verify, db_verify):
        heading = doc.add_heading(title_text, level=2)
        
        # Script part
        p_script_label = doc.add_paragraph()
        p_script_label.add_run('[Script]').bold = True
        p_script_label.add_run(':')
        
        doc.add_paragraph(script_text)
        
        # Verify part
        p_verify_label = doc.add_paragraph()
        p_verify_label.add_run('[How to Verify]').bold = True
        p_verify_label.add_run(':')
        
        if ui_verify:
            p_ui = doc.add_paragraph(style='List Bullet')
            p_ui.add_run('UI Validation: ').bold = True
            p_ui.add_run(ui_verify)
            
        if db_verify:
            p_db = doc.add_paragraph(style='List Bullet')
            p_db.add_run('Database Validation: ').bold = True
            p_db.add_run(db_verify)
            
        doc.add_paragraph()

    # Opening
    doc.add_heading('Opening', level=2)
    doc.add_paragraph('[Script]:').runs[0].bold = True
    doc.add_paragraph("Welcome everyone. In this video, I will walk you through the comprehensive progress we've made on Mamba Phase 2. We have hit major milestones across our multi-tenant architecture, data security, and core platform features. I want to show you exactly what has been built, how it works, and how you can verify it directly in our environments.")
    doc.add_paragraph()

    # Sections
    add_section(
        '1. Multi-Tenant Architecture & Data Isolation',
        "We started with the foundational piece: true multi-tenancy. Agencies and Sellers are now fully isolated. An agency can invite and link sellers, and account managers are strictly scoped to only see data for the sellers assigned to them. We enforce this directly at the database level using Row Level Security that keys off the authentication context.",
        "Log in as an Agency Admin, invite a seller, and accept the invite. Assign an Account Manager to that seller. Log in as that Account Manager and verify that only that specific seller's data is visible. Unlink the seller and ensure access is instantly revoked.",
        "Run a query on the public.tenants table to check that the parent_tenant_id correctly points to the agency. Check public.user_seller_assignments to confirm strict assignments."
    )

    add_section(
        '2. Role-Based Access Control (RBAC) & Custom Roles',
        "Next, we deployed a robust RBAC system. System roles are seeded, and we implemented custom roles that belong to specific tenants. Permissions are resolved securely at the database level, which prevents any API tampering or bypass attempts.",
        "As an Agency Admin, open team management and create a custom role with specific toggles. Assign it to a user, log in as that user, and verify access matches exactly.",
        "Query public.get_user_effective_permissions_on_tenant to see the resolved permissions matrix for a given user UUID."
    )

    add_section(
        '3. Restricted Financial Visibility',
        "To protect sensitive data, we built seller-controlled financial visibility. Sellers can hide specific financial metrics, like COGS and margins, from their agencies. The seller restrictions override all agency permissions.",
        "As a Seller Admin, restrict COGS visibility for the linked agency. Log in as an agency user and check dashboards or PDF exports—the restricted numbers will be masked as 'Restricted'.",
        "Query public.seller_financial_visibility_rules and use the public.get_financial_field_access function to verify that can_view_cogs returns false for the agency user."
    )

    add_section(
        '4. Production Infrastructure & Audit Logging',
        "We hardened our production infrastructure with comprehensive audit logging. Any sensitive mutation—like role changes or financial visibility updates—is immutably recorded with before and after states.",
        None,
        "Query public.audit_logs filtering for the last 7 days to see a detailed history of actions like financial.restriction_change and tenant.link_created."
    )

    add_section(
        '5. White-Label Customization',
        "We delivered White-Label Platform Customization, allowing agencies to brand the platform as their own. The branding is owned and configured at the Agency tenant level.",
        "As an Agency Admin, set a custom logo, primary color, and display name. Log in as a seller under that agency and verify the UI updates to reflect the custom agency palette and branding.",
        "Check the public.tenant_branding table for the agency's UUID and the public.tenant_branding_audit table for the history of branding changes."
    )

    add_section(
        '6. TikTok Data Ingestion Architecture',
        "Our TikTok data ingestion architecture is fully operational, complete with health tracking and refresh mechanisms. We also added an ingestion monitoring console for platform operators.",
        "As a Seller Admin, connect TikTok OAuth and trigger a sync. As a platform super admin, open the Ingestion Monitoring console to track sync status and catch any errors.",
        "Check public.ingestion_jobs and public.tiktok_shops to verify token statuses and sync job outcomes."
    )

    add_section(
        '7. Unified Messaging Inbox (Email-First)',
        "A major recent addition is the Unified Messaging Inbox. It uses email as the transport layer but presents a seamless, chat-style interface inside Mamba, maintaining strict tenant boundaries.",
        "Open the Messaging interface, select a seller context, and start a thread. Send a reply from an external mailbox and watch it land in the same thread within the app. Log in as an unassigned agency user and verify the thread is hidden.",
        "Query public.messaging_conversations and public.messaging_messages to confirm threads are correctly persisted and isolated by seller_tenant_id."
    )

    add_section(
        '8. Dashboard Export & Automated Email Delivery',
        "We built PDF dashboard exports and recurring email schedules, allowing automated reporting directly to clients.",
        "From the dashboard, trigger a one-off email export and confirm delivery. Then, create a recurring schedule and verify it appears in the active schedule list.",
        "Check public.dashboard_email_schedules to confirm the schedule data and review public.audit_logs for export.dashboard_email events."
    )

    add_section(
        '9. Customer Bug Reporting & Ticketing Integration',
        "Finally, we integrated a seamless in-app bug reporting and ticketing system that hooks directly into our support workflows.",
        "Submit a bug report via the in-app Support widget. Confirm the success message and verify the ticket lands in the configured support inbox.",
        "Run a query on public.support_ticket_submissions to verify the submission record, including vendor, external IDs, and tenant context."
    )

    # Closing
    doc.add_heading('Closing', level=2)
    doc.add_paragraph('[Script]:').runs[0].bold = True
    doc.add_paragraph("That covers the major technical milestones for Mamba Phase 2. We established a secure, scalable foundation and delivered high-impact features that drive immediate value. Thank you for watching, and let me know if you have any questions.")

    doc.save('/Users/David/Documents/Mamba/Mamba_Phase_2_Progress_Video_Script.docx')

if __name__ == '__main__':
    create_script()
