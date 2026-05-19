#!/usr/bin/env python3
"""Generate client-facing security & UX remediation DOCX."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).resolve().parent.parent / "docs" / "client-deliverables"
OUT_FILE = OUT_DIR / "Mamba_Security_and_UX_Remediation_Report.docx"


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def add_issue_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    headers = ["Issue reported", "Root cause", "Remediation"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1A3D38")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
                r.font.name = "Calibri"
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Calibri"
    return table


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Mamba Platform\nSecurity & UX Remediation Report")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.name = "Calibri"
    tr.font.color.rgb = RGBColor(26, 61, 56)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Prepared for client review  |  {date.today().strftime('%B %d, %Y')}")
    sr.font.size = Pt(11)
    sr.font.name = "Calibri"
    sr.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    add_heading(doc, "Executive summary", 1)
    add_para(
        doc,
        "This document summarizes issues identified during the security and usability review of the Mamba "
        "platform, along with the remediations implemented in the application codebase and database migrations. "
        "All items listed below have been addressed. Where an item was determined to be expected behavior rather "
        "than a defect, that clarification is documented explicitly.",
    )

    add_heading(doc, "Deployment note", 2)
    add_bullets(
        doc,
        [
            "Database changes are delivered as Supabase SQL migrations in the repository and must be applied to each environment (staging, production) before full effect.",
            "Backend (Express API) and frontend (Vite) deployments should be updated together with migrations.",
            "Key migrations for security items: 20260515180000, 20260515190000, 20260515200000, 20260515210000, 20260515220000.",
        ],
    )

    # --- UI/UX ---
    add_heading(doc, "1. UI / UX remediations", 1)

    ui_rows = [
        (
            "Team invite email sent twice; only the second email had a valid accept link.",
            "The invite API called Supabase Auth invite/OTP in addition to a custom HTML email, producing duplicate messages with inconsistent links.",
            "Consolidated to a single outbound email via the application mailer with one accept-invitation deep link. New users are created without triggering Supabase’s duplicate invite mailer.",
        ),
        (
            "Invite email should include the organization name.",
            "Template did not consistently pass tenant display name.",
            "Invite emails now include the inviting organization’s name in the subject and body (buildTeamInviteEmailHtml).",
        ),
        (
            "Overview key metrics: “down” trend icon incorrectly oriented.",
            "Negative trend reused TrendingUp with CSS rotation instead of a proper down icon.",
            "Replaced with TrendingDown for negative trends; active Today/Yesterday states use consistent brand styling.",
        ),
        (
            "P&L page should default to Yesterday.",
            "Default date preset was last-30-days.",
            "Default preset changed to yesterday (shop timezone). Today/Yesterday controls and export metadata aligned with settlement vs paid-time semantics.",
        ),
        (
            "Seller Admin Home: “Connect new shop” does nothing.",
            "Wrong account was resolved (first account in list); write-access RPC had a tenant-join bug; errors were not surfaced in the UI.",
            "Account resolution now uses the seller’s canonical account; write permission checks corrected; errors displayed; button hidden when the user lacks TikTok connect permission.",
        ),
    ]
    add_issue_table(doc, ui_rows)

    # --- Security RLS ---
    add_heading(doc, "2. Security — cross-tenant data exposure (accounts & shops)", 1)
    add_para(
        doc,
        "Reported behavior: GET /rest/v1/accounts and GET /rest/v1/tiktok_shops returned rows belonging to "
        "tenants outside the authenticated user’s membership.",
    )
    add_para(doc, "Remediation:", bold=True)
    add_bullets(
        doc,
        [
            "Restored and hardened Row Level Security (RLS) on accounts and tiktok_shops using account_is_visible_to_user and tenant_is_visible_to_user (memberships, agency links, AM/AC assignments, platform super admin).",
            "Removed legacy permissive policies that could allow broader reads when combined with stale helpers.",
            "Aligned server-side check_user_account_access RPC with the same visibility rules used by RLS.",
            "Added defense-in-depth tenant filtering on the home shop list in the frontend.",
            "Migration: 20260515190000_harden_accounts_shops_rls.sql",
        ],
    )

    # --- Profiles ---
    add_heading(doc, "3. Profiles endpoint — role semantics and duplicate requests", 1)
    add_para(
        doc,
        'Reported: profiles.role returns "client"; profile request fires three times on load.',
    )
    add_para(doc, "Clarification and remediation:", bold=True)
    add_bullets(
        doc,
        [
            'profiles.role is a legacy account-type flag. The value "client" is normal for seller and agency users; it does not mean “client portal only.” Actual permissions come from tenant_memberships and RBAC (e.g. Seller Admin, Agency Admin).',
            "Platform operators may still use profiles.role = admin or Super Admin membership.",
            "Profile loading deduplicated via React Query (single cached fetch per user session) in AuthContext; TenantContext reuses that data instead of issuing redundant full profile queries.",
            "Misleading “Account type: client” copy removed from Profile UI; developer documentation added in profileApi.ts and supabase.ts.",
        ],
    )

    # --- parent_tenant_id ---
    add_heading(doc, "4. Tenant model — parent_tenant_id", 1)
    add_para(
        doc,
        "Reported: unclear purpose of tenants.parent_tenant_id.",
    )
    add_para(doc, "Clarification (documented in database and code):", bold=True)
    add_bullets(
        doc,
        [
            "parent_tenant_id links a seller tenant to its parent agency when an agency–seller relationship is accepted.",
            "It is a single-level agency→seller reference only — not permission inheritance, not a multi-level reseller tree, and not used for arbitrary tenant hierarchies.",
            "SQL comments added in migration 20260515200000_document_tenant_parent_tenant_id.sql; developer reference in src/lib/tenantModel.ts.",
        ],
    )

    # --- user_can_access_account ---
    add_heading(doc, "5. Access control — user_can_access_account RPC", 1)
    add_para(
        doc,
        "Reported concern: frontend RPC appears to be the security boundary for shop access.",
    )
    add_para(doc, "Clarification and documentation:", bold=True)
    add_bullets(
        doc,
        [
            "user_can_access_account is a UI convenience only (ShopPage gate, button visibility). It uses SECURITY INVOKER and mirrors account_is_visible_to_user.",
            "Authoritative enforcement: RLS on all shop-scoped tables for direct PostgREST access; Express middleware (verifyAccountIdParam, enforceRequestAccountAccess) calling check_user_account_access / check_user_account_write_access (service role only) for /api/* routes.",
            "Documented in src/lib/accessControl.ts and migration 20260515210000_document_account_access_security_model.sql.",
            "Outstanding hardening note: /api/tiktok-shop/debug/* should receive the same account or admin auth if retained in production.",
        ],
    )

    # --- Post logout ---
    add_heading(doc, "6. Post-logout token replay", 1)
    add_para(
        doc,
        "Reported: copied cURL requests with a bearer token still succeed after logout.",
    )
    add_para(doc, "Root cause:", bold=True)
    add_para(
        doc,
        "JWT access tokens remain cryptographically valid until expiry unless the server verifies that the "
        "session row still exists in auth.sessions. Logout previously cleared client storage but did not "
        "invalidate in-flight access tokens for API/PostgREST replay.",
    )
    add_para(doc, "Remediation:", bold=True)
    add_bullets(
        doc,
        [
            "Logout now uses signOut({ scope: 'global' }) to revoke refresh tokens and delete session rows server-side.",
            "PostgREST pre-request hook validate_active_auth_session rejects authenticated requests whose session_id is no longer in auth.sessions (migration 20260515220000_validate_active_auth_session.sql).",
            "Express API validates bearer tokens via resolveUserIdFromBearerToken: auth.getUser plus auth_session_is_active RPC.",
            "Recommendation: keep JWT expiry at one hour or less in Supabase Auth settings.",
        ],
    )

    # --- Seller admin invite ---
    add_heading(doc, "7. Permissions — Seller Admin team invites", 1)
    add_para(
        doc,
        "Reported: Seller Admins cannot invite users to their seller account.",
    )
    add_para(doc, "Finding:", bold=True)
    add_para(
        doc,
        "This is expected product behavior, not a permissions defect, provided the user holds the Seller Admin "
        "system role (or a custom role with users.manage) on the seller tenant.",
    )
    add_bullets(
        doc,
        [
            "Path: Console → Team & roles → Assign role to user → Invite by Email (or search existing user).",
            "Backend: POST /api/team/invite-member guarded by user_can_manage_tenant_members, which allows Seller Admin and parent Agency Admin on seller tenants.",
            "Users with Seller User only cannot manage team membership by design.",
            "If a specific user cannot access Team & roles, verify their tenant_memberships role is Seller Admin (active), not Seller User.",
        ],
    )

    # --- Verification ---
    add_heading(doc, "8. Suggested verification checklist", 1)
    verify_rows = [
        ("Duplicate invite email", "Invite a new email; confirm exactly one email with working /accept-invitation link and org name in body."),
        ("RLS / accounts & shops", "As User A, call accounts/tiktok_shops REST; confirm only visible tenants/shops are returned."),
        ("Post-logout replay", "Copy authenticated cURL, log out globally, replay; expect 401 or session invalid error after migration applied."),
        ("Connect new shop", "As Seller Admin with tiktok.auth, connect flow completes or shows explicit error; button hidden without permission."),
        ("P&L default", "Open P&L; confirm Yesterday is selected by default."),
        ("Seller Admin invite", "Console → Team & roles → Assign role to user → Invite by Email succeeds for Seller Admin."),
    ]
    add_issue_table(doc, verify_rows)

    add_heading(doc, "9. Contact & follow-up", 1)
    add_para(
        doc,
        "If any item fails verification after migrations and application deploy, please capture the HTTP "
        "status code, endpoint, user role, and tenant context so engineering can reproduce quickly.",
    )

    doc.save(OUT_FILE)
    print(f"Wrote {OUT_FILE}")


if __name__ == "__main__":
    build()
